from __future__ import annotations

import datetime
import os
from typing import Optional

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from .feature_catalog import FeatureCatalog


# High-level, non‑technical capabilities per domain, written for school admins.
DOMAIN_CAPABILITIES = {
    "Student Information Management": {
        "summary": "Central place to store and view every student's profile, history, and key details.",
        "bullets": [
            "Keep complete student profiles (contact details, emergency contacts, background).",
            "See enrollment history, current class, and status for each student.",
            "Search and filter students by grade, class, or other criteria.",
            "Support better decisions with a single, up‑to‑date source of student information.",
        ],
    },
    "Class & Timetable Management": {
        "summary": "Tools to organize classes, timetables, and groupings across the school.",
        "bullets": [
            "Set up classes and sections for each grade.",
            "Assign teachers and students to classes.",
            "Manage changes in class structure during the school year.",
            "Give staff a clear view of who is teaching what, and where.",
        ],
    },
    "Attendance Management": {
        "summary": "End‑to‑end tracking of student attendance for each class and school day.",
        "bullets": [
            "Record daily attendance and late arrivals quickly.",
            "View attendance by student, class, grade, or whole school.",
            "Identify patterns of absence early, for faster follow‑up.",
            "Generate attendance reports to share with teachers and parents.",
        ],
    },
    "Examinations & Assessment": {
        "summary": "Plan, record, and review student assessments and exam results.",
        "bullets": [
            "Capture scores for exams, quizzes, and other assessments.",
            "Track performance over time for each student and class.",
            "Support grading policies and calculations.",
            "Provide a clearer academic picture for staff and families.",
        ],
    },
    "Finance & Fees": {
        "summary": "Monitor fee collection and financial interactions with students and families.",
        "bullets": [
            "Record fee structures, discounts, and payment plans.",
            "Track payments and outstanding balances for each student.",
            "Support clear reporting on collections and arrears.",
            "Help leadership see the financial health of the school.",
        ],
    },
    "Grades & Academic Performance": {
        "summary": "Monitor how students are performing across subjects and terms.",
        "bullets": [
            "Store grades for each subject and term.",
            "View performance by student, class, grade level, or subject.",
            "Identify top performers and learners who need extra support.",
            "Provide data to guide academic interventions and planning.",
        ],
    },
    "Teacher Portal": {
        "summary": "A dedicated workspace for teachers to manage their daily tasks.",
        "bullets": [
            "Give teachers a clear view of their classes and students.",
            "Make it easy to record attendance and grades.",
            "Provide quick access to key student information.",
            "Reduce paperwork so teachers can focus more on teaching.",
        ],
    },
    "Parent Portal": {
        "summary": "Online access for parents to see key information about their children.",
        "bullets": [
            "Allow parents to view attendance, grades, and announcements.",
            "Improve communication between home and school.",
            "Increase transparency and trust with families.",
            "Reduce the need for manual updates and paper reports.",
        ],
    },
    "Administration & Configuration": {
        "summary": "High‑level tools for school leaders and system administrators.",
        "bullets": [
            "Configure core settings, academic years, and structures.",
            "Manage user accounts and access to the system.",
            "Monitor key activity and ensure data quality.",
            "Support long‑term planning and policy decisions.",
        ],
    },
    "Advanced Administration": {
        "summary": "Extra‑granular controls for complex or multi‑school environments.",
        "bullets": [
            "Manage permissions and roles at a very detailed level.",
            "Configure advanced security or audit requirements.",
            "Support complex, multi‑campus or multi‑school setups.",
        ],
    },
    "CRM & Customer Management": {
        "summary": "Tools to manage relationships with external partners, customers, or stakeholders.",
        "bullets": [
            "Track interactions with partners, customers, or donors.",
            "Organize tasks, tickets, and follow‑ups.",
            "Support communication and long‑term relationship building.",
        ],
    },
    "Authentication, Security & Access Control": {
        "summary": "Keep the system secure and ensure each person only sees what they should.",
        "bullets": [
            "Protect access with user accounts and secure sign‑in.",
            "Control which modules each role (admin, teacher, parent, etc.) can use.",
            "Reduce the risk of data breaches and misuse.",
        ],
    },
    "Localization & Multi‑Language Support": {
        "summary": "Make the system comfortable for users in different languages.",
        "bullets": [
            "Offer the interface in multiple languages.",
            "Support schools and users from different regions.",
            "Reduce confusion by letting users work in their preferred language.",
        ],
    },
}


def _set_document_styles(document: Document) -> None:
    styles = document.styles

    # Base font
    try:
        normal_style = styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = "Calibri"
        normal_font.size = Pt(11)
    except KeyError:
        # Fallback silently if the style is not present
        pass

    # Headings – use dict-style access to be compatible with python-docx
    for level in range(1, 4):
        style_name = f"Heading {level}"
        try:
            style = styles[style_name]
        except KeyError:
            continue
        style.font.name = "Segoe UI"
        style.font.bold = True


def _add_title_page(document: Document, system_name: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.5)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Title
    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run(system_name)
    run.bold = True
    run.font.size = Pt(28)

    document.add_paragraph()

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run("System Features & Technical Overview")
    subtitle_run.font.size = Pt(16)

    document.add_paragraph()

    today = datetime.date.today().strftime("%B %d, %Y")
    date_p = document.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_p.add_run(f"Generated on {today}")
    date_run.font.size = Pt(11)

    # Company placeholder
    document.add_paragraph()
    company_p = document.add_paragraph()
    company_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_p.add_run("Prepared by: ______________________________")

    # New section for content
    document.add_section(WD_SECTION_START.NEW_PAGE)


def _add_executive_summary(document: Document, catalog: FeatureCatalog) -> None:
    document.add_heading("Executive Summary", level=1)

    summary = catalog.infer_high_level_summary()
    document.add_paragraph(
        summary
        + " The system is designed to help school leaders, principals, and "
        "administrators see the full picture of their institution in one place."
    )

    if catalog.domains:
        p = document.add_paragraph()
        p.add_run("At a glance, this MIS helps you to:").bold = True
        for domain in catalog.all_domains_sorted():
            document.add_paragraph(f"Manage {domain.name.lower()}.", style="List Bullet")


def _add_domain_sections(document: Document, catalog: FeatureCatalog) -> None:
    document.add_heading("System Domains & Features", level=1)

    for domain in catalog.all_domains_sorted():
        document.add_heading(domain.name, level=2)

        caps = DOMAIN_CAPABILITIES.get(domain.name)
        if caps and caps.get("summary"):
            document.add_paragraph(caps["summary"])
        else:
            document.add_paragraph(
                f"This area of the system contains the main tools your staff use to "
                f"handle {domain.name.lower()}."
            )

        bullets = (caps or {}).get("bullets") or []
        if bullets:
            for text in bullets:
                document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(
                f"Key actions in this area include viewing information, updating "
                f"records, and producing reports related to {domain.name.lower()}.",
                style="List Bullet",
            )


def _add_security_section(document: Document, catalog: FeatureCatalog) -> None:
    if "Authentication, Security & Access Control" not in catalog.domains:
        return

    document.add_heading("Security, Authentication & Access Control", level=1)
    document.add_paragraph(
        "The MIS includes built‑in authentication and role‑based access control "
        "mechanisms. Dedicated authentication screens, contexts, and access control "
        "components are used to manage sessions, user identities, and permissions."
    )
    document.add_paragraph(
        "Role‑aware components and context providers ensure that sensitive data and "
        "administrative capabilities are only available to authorized users."
    )


def _add_localization_section(document: Document, catalog: FeatureCatalog) -> None:
    locales = catalog.get_supported_locales()
    if not locales:
        return

    document.add_heading("Localization & Multi‑Language Support", level=1)

    p = document.add_paragraph()
    p.add_run("Supported Locales. ").bold = True
    p.add_run(
        "The MIS ships with localization resources for the following locales, "
        "making it suitable for multilingual deployments:"
    )

    for loc in locales:
        document.add_paragraph(loc, style="List Bullet")

    document.add_paragraph(
        "Localization keys and translation files are centrally managed, allowing "
        "new languages to be added with minimal changes to the application code."
    )


def build_docx(
    catalog: FeatureCatalog,
    project_root: str,
    output_path: str,
    system_name: Optional[str] = None,
) -> str:
    """
    Build the DOCX document and return the absolute path to the generated file.
    """
    if not system_name:
        system_name = "School Management Information System (MIS)"

    document = Document()
    _set_document_styles(document)
    _add_title_page(document, system_name=system_name)
    _add_executive_summary(document, catalog)
    _add_domain_sections(document, catalog)
    _add_security_section(document, catalog)
    _add_localization_section(document, catalog)

    output_abs = os.path.abspath(output_path)
    os.makedirs(os.path.dirname(output_abs), exist_ok=True)
    document.save(output_abs)
    return output_abs


